import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

# 输入输出路径
input_path = "data/7.xlsx"  # 请改成你的源文件路径
output_dir = "data_7"            # 输出目录
os.makedirs(output_dir, exist_ok=True)

# 读取Excel数据，假设第一行为标题
df = pd.read_excel(input_path, header=1)

# 需要保留的列名（根据你给的目标格式）
cols = ['网格名称', '2020年', '2024 年', '2025 年', '2026 年', '2027 年', '2028 年', '2029 年', '2030年', '“十四五”增速', '“十五五”增速']

# 清理列名空格，方便操作
df.columns = [col.strip().replace(' ', '') for col in df.columns]

# 重新映射你需要的列名对应关系，方便后续写表头（统一无空格）
col_map = {
    '网格名称': '网格名称',
    '2020年': '2020年',
    '2024年': '2024年',
    '2025年': '2025年',
    '2026年': '2026年',
    '2027年': '2027年',
    '2028年': '2028年',
    '2029年': '2029年',
    '2030年': '2030年',
    '“十四五”增速': '十四五增速',
    '“十五五”增速': '十五五增速'
}

# 选取对应列，注意2024年等列名去空格后可能是'2024年'
selected_cols = [col.strip().replace(' ', '') for col in cols]
df_selected = df[selected_cols]

# 按网格名称分组并分别保存
for grid_name, group_df in df_selected.groupby('网格名称'):
    # 创建Word文档
    doc = Document()
    # 设置字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    doc.add_heading(f"表7  {grid_name}全社会用电负荷预测表    单位：MW ，%", level=1)

    # 新建表格，1行表头 + 1行数据，列数与选列数对应
    table = doc.add_table(rows=2, cols=len(selected_cols))
    table.style = 'Table Grid'

    # 写入表头
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(selected_cols):
        hdr_cells[i].text = col_map.get(col, col)

    # 写入数据（该网格名只有一条记录）
    row_cells = table.rows[1].cells
    row = group_df.iloc[0]
    for i, col in enumerate(selected_cols):
        val = row[col]
        # 格式化浮点数保留两位小数
        if pd.isna(val):
            row_cells[i].text = '/'
        elif isinstance(val, float):
            row_cells[i].text = f"{val:.2f}"
        else:
            row_cells[i].text = str(val)

    # 保存文件名安全处理
    safe_name = "".join(c for c in grid_name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    doc.save(f"{output_dir}/{safe_name}.docx")

print("所有网格文件已生成完毕！")