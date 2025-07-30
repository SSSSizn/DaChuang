import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_vmerge(cell, merge_type):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    vMerge.set(qn('w:val'), merge_type)
    tcPr.append(vMerge)

def set_cell_hmerge(cell, merge_type):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    hMerge = OxmlElement('w:hMerge')
    hMerge.set(qn('w:val'), merge_type)
    tcPr.append(hMerge)

excel_path = "data/17.xlsx"  # 修改为你的文件路径
output_dir = "data_17"
os.makedirs(output_dir, exist_ok=True)

df = pd.read_excel(excel_path)

total_cols = 25  # 0-24共25列，调整根据实际情况
# 这里默认DataFrame列顺序是和Word列一致的，如果不一致，需调整下方取值方式

for grid_name, group_df in df.groupby('所属网格'):
    doc = Document()
    doc.add_heading(f"{grid_name} 项目建设统计表", level=1)

    table = doc.add_table(rows=2 + len(group_df), cols=total_cols)
    table.style = 'Table Grid'

    # 表头第一行
    first_row = [
        "序号", "所属单位", "所属网格", "项目名称", "建设年份", "电压等级", "项目类型", "项目场景",
        "项目属性1", "项目属性2", "项目属性3", "建设规模（座，台，km，kVA，套）"
    ] + [""]*11 + [  # 建设规模跨12列
        "项目投资（万元）", "对应解决负面问题/需求清单"
    ]

    # 表头第二行
    second_row = [""]*11 + [
        "架空线", "电缆线", "环网箱", "环网室", "配变", "配变容量", "柱上开关", "DTU", "FTU", "智能融合终端", "光缆", "ONU(套)"
    ] + [""]*2

    # 填写第一行
    cells = table.rows[0].cells
    for i, text in enumerate(first_row):
        cells[i].text = text

    # 填写第二行
    cells = table.rows[1].cells
    for i, text in enumerate(second_row):
        cells[i].text = text

    # 跨两行单元格（前11列）
    for i in range(11):
        set_cell_vmerge(table.cell(0, i), 'restart')
        set_cell_vmerge(table.cell(1, i), 'continue')

    # 建设规模12列横向合并
    set_cell_hmerge(table.cell(0, 11), 'restart')
    for i in range(1, 12):
        set_cell_hmerge(table.cell(0, 11 + i), 'continue')

    # 项目投资和需求清单跨两行
    for i in range(23, 25):
        set_cell_vmerge(table.cell(0, i), 'restart')
        set_cell_vmerge(table.cell(1, i), 'continue')

    # 填充数据行
    for row_idx, row in enumerate(group_df.itertuples(index=False), start=2):
        cells = table.rows[row_idx].cells
        cells[0].text = str(row[0])  # 序号
        cells[1].text = str(row[1])  # 所属单位
        cells[2].text = str(row[2])  # 所属网格
        cells[3].text = str(row[3])  # 项目名称
        cells[4].text = str(row[4])  # 建设年份
        cells[5].text = str(row[5])  # 电压等级
        cells[6].text = str(row[6])  # 项目类型
        cells[7].text = str(row[7])  # 项目场景
        cells[8].text = str(row[8])  # 项目属性1
        cells[9].text = str(row[9])  # 项目属性2
        cells[10].text = str(row[10])  # 项目属性3

        # 建设规模12列
        for i, col_idx in enumerate(range(11, 23)):
            val = getattr(row, f'_ {col_idx + 1}') if f'_ {col_idx + 1}' in row._fields else ''
            if col_idx < len(row):
                val = row[col_idx]
            else:
                val = ''
            cells[col_idx].text = str(val) if pd.notna(val) else ''

        # 项目投资（万元）
        cells[23].text = str(row[23]) if 23 < len(row) and pd.notna(row[23]) else ''

        # 对应解决负面问题/需求清单
        cells[24].text = str(row[24]) if 24 < len(row) and pd.notna(row[24]) else ''

    safe_name = grid_name.replace('/', '_').replace('\\', '_').replace(' ', '_')
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))
    print(f"已生成: {safe_name}.docx")
