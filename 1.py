import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

excel_path = "data/1.xlsx"

df = pd.read_excel(excel_path, skiprows=1)

df_selected = df[[
    '网格名称',
    '规划特征',
    '供电面积（km2）',
    '10（20）kV线路长度架空长度（km）',
    '10（20）kV线路长度电缆长度（km）',
    '10（20）kV线路公变配电容量（MVA）',
    '10（20）kV线路专变配电容量（MVA）',
    '低压用户数（万户）'
]].copy()

df_selected.columns = ['名称', '类型', '供电面积', '架空', '电缆', '公变', '专变', '低压用户数']

groups = df_selected.groupby('名称')

for name, group_df in groups:
    doc = Document()

    p = doc.add_heading(f'表1  {name}概况   单位：km2 ，MVA，万户', level=1)
    run = p.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 先创建两行表头（多级表头）
    table = doc.add_table(rows=2, cols=8)
    table.style = 'Table Grid'
    table.autofit = True

    # 第一行：主表头（合并单元格）
    header1 = table.rows[0].cells
    header1[0].text = '名称'
    header1[1].text = '类型'
    header1[2].text = '供电面积'
    header1[3].text = '10kV线路长度'
    header1[5].text = '10kV配变容量'
    header1[7].text = '低压用户数'

    # 合并跨列单元格
    header1[3].merge(header1[4])  # 合并"架空" + "电缆"
    header1[5].merge(header1[6])  # 合并"公变" + "专变"

    # 合并竖向单元格
    header1[0].merge(table.cell(1, 0))
    header1[1].merge(table.cell(1, 1))
    header1[2].merge(table.cell(1, 2))
    header1[7].merge(table.cell(1, 7))

    # 第二行：子表头
    header2 = table.rows[1].cells
    header2[3].text = '架空'
    header2[4].text = '电缆'
    header2[5].text = '公变'
    header2[6].text = '专变'

    # 添加数据行
    for idx, row in group_df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = '' if pd.isna(val) else str(val)

    # 文件名处理
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    output_word = f"grid_1/{safe_name}.docx"
    doc.save(output_word)
    print(f"已保存：{output_word}")
