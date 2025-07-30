import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ======== 配置路径 ==========
input_path = "data/11.xlsx"  # 源文件路径
output_dir = "data_11"  # 输出目录
os.makedirs(output_dir, exist_ok=True)

# ======== 读取数据 ==========
df = pd.read_excel(input_path, header=1)

df["电压等级"] = df["电压等级"].str.strip()
df["类型"] = df["类型"].str.strip()
df["网格名称"] = df["网格名称"].str.strip()

years = ["2024年", "2025年", "2026年", "2027年", "2028年", "2029年", "2030年"]
group_columns = ["电压等级", "类型"]

# 获取所有网格
for grid_name, group_df in df.groupby("网格名称"):
    # 汇总该网格下的电压等级-类型组合
    summary = group_df.groupby(group_columns)[years].sum().reset_index()

    # 创建 Word 文档
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 添加标题
    doc.add_heading(grid_name, level=1)

    # 表格：列数 = 电压等级 + 类型 + 年份（无细化指标）
    table = doc.add_table(rows=1, cols=2 + len(years))
    table.style = 'Table Grid'

    # ===== 写表头 =====
    header = table.rows[0].cells
    header[0].text = '电压等级（kV）'
    header[1].text = '类型'
    for i, y in enumerate(years):
        header[i + 2].text = y

    # ===== 写数据 =====
    prev_voltage = None  # 用于控制合并格式
    for _, row in summary.iterrows():
        cells = table.add_row().cells
        voltage = row["电压等级"]
        cells[0].text = voltage if voltage != prev_voltage else ""
        cells[1].text = row["类型"]
        for i, y in enumerate(years):
            value = row[y]
            cells[i + 2].text = str(round(value, 2)) if pd.notna(value) else '\\'
        prev_voltage = voltage

    # 保存
    safe_name = grid_name.replace("/", "-").replace("\\", "-")
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))

print("✔ 所有网格储能装机数据已导出完毕！")
