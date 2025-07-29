import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import os
import re

excel_path = "data/3.xlsx"
os.makedirs("data_3", exist_ok=True)
df = pd.read_excel(excel_path, skiprows=1)

# 字符清洗函数
def normalize(s):
    s = str(s).strip()
    s = s.replace("（", "(").replace("）", ")")  # 全角转半角括号
    s = s.replace("\xa0", "")  # 非断行空格（U+00A0）
    s = s.replace(" ", "")  # 去普通空格
    s = re.sub(r"\s+", "", s)  # 去掉所有空白字符
    return s

def normalize_tuple(t):
    return tuple(normalize(x) for x in t)

def set_font(cell, text):
    run = cell.paragraphs[0].add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

df_selected = df[[
    '网格名称',
    '项目',
    '2020年',
    '2021年',
    '2022年',
    '2023年',
    '2024年',
]].copy()

groups = df.groupby('网格名称')

for name, group_df in groups:
    group_df = group_df.copy()

    # 提取电量和负荷两类数据（注意类型字段需要提前清洗）
    group_df['项目'] = group_df['项目'].apply(normalize)
    filtered = group_df[group_df['项目'].isin(['电量', '负荷'])]

    # 如果没有数据就跳过
    if filtered.empty:
        print(f"{name} 无有效数据，跳过")
        continue

    # 创建 Word 文档
    doc = Document()
    p = doc.add_heading(f'表 3  {name}全社会负荷及电量规模    单位：MW，亿kWh', level=1)
    run = p.runs[0]
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    run.bold = True

    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    headers = ['项目', '2020年', '2021年', '2022年', '2023年', '2024年']
    for i, h in enumerate(headers):
        set_font(table.cell(0, i), h)

    # 添加两行数据
    for item in ['负荷','电量']:
        row = table.add_row().cells
        row[0].text = item
        row_data = filtered[filtered['项目'] == item]
        if not row_data.empty:
            row_values = row_data.iloc[0][['2020年', '2021年', '2022年', '2023年', '2024年']]
            for j, v in enumerate(row_values):
                row[j + 1].text = '0' if pd.isna(v) else str(v)
        else:
            for j in range(1, 6):
                row[j].text = '/'

    # 保存文档
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    output_word = f"data_3/{safe_name}.docx"
    doc.save(output_word)
    print(f"已保存：{output_word}")
