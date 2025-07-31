import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

input_path = "data/5.xlsx"  # 源文件路径
output_dir = "data_5"
os.makedirs(output_dir, exist_ok=True)

# 读取数据
df = pd.read_excel(input_path, header=1)

# 年份列
years = ['2025年容量（MVA）', '2026年容量（MVA）', '2027年容量（MVA）', '2028年容量（MVA）', '2029年容量（MVA）', '2030年容量（MVA）']

# 按网格名称分组
groups = df.groupby('网格名称')

def set_font(cell, text):
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

for grid_name, group in groups:
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_heading(f"表5 {grid_name} “十五五”110kV 变电站规划建设情况  单位：座、台、 MVA", level=1)

    cols = 3 + len(years)
    # 只创建1行表头
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'

    # 写表头
    hdr_cells = table.rows[0].cells
    set_font(hdr_cells[0], '序号')
    set_font(hdr_cells[1], '变电站名称')
    set_font(hdr_cells[2], '分类')
    for i, y in enumerate(years):
        # 去掉“容量（MVA）”部分，只保留年份数字
        set_font(hdr_cells[i+3], y.replace('容量（MVA）',''))

    # 写数据，每个变电站两行（主变台数 + 主变容量）
    for idx, (_, row) in enumerate(group.iterrows(), 1):
        # 主变台数行
        cells1 = table.add_row().cells
        set_font(cells1[0], str(idx))
        set_font(cells1[1], row['变电站名称'])
        set_font(cells1[2], '主变台数')
        for i, y in enumerate(years):
            val = str(row['10（20）kV间隔数量']) if i == 0 and pd.notna(row['10（20）kV间隔数量']) else '/'
            set_font(cells1[i+3], val)

        # 主变容量行
        cells2 = table.add_row().cells
        set_font(cells2[0], '')
        set_font(cells2[1], '')
        set_font(cells2[2], '主变容量')
        for i, y in enumerate(years):
            val = str(row[y]) if pd.notna(row[y]) else '/'
            set_font(cells2[i+3], val)

    # 保存文件，文件名去除非法字符
    safe_name = "".join(c for c in grid_name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))

print("所有网格变电站规划数据已导出完毕！")