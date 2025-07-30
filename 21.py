import pandas as pd
from docx import Document
from collections import defaultdict
import os

# === 文件路径 ===
excel_path = "data/21.xlsx"
output_dir = "data_21"
os.makedirs(output_dir, exist_ok=True)

# === 读取 Excel，三层表头 ===
df = pd.read_excel(excel_path, header=[0, 1, 2])

# === 拆分元数据与数据 ===
df_meta = df.iloc[:, :2]
df_data = df.iloc[:, 2:]

# === 类型顺序 ===
type_order = [
    '中压新建项目', '中压改造项目',
    '低压新建项目', '低压改造项目',
    '智能化项目', '其他项目'
]

# === 每个网格一份 Word ===
for idx, row in df.iterrows():
    grid_code = row.iloc[0]
    grid_name = row.iloc[1]

    # 构造类型 -> 项目 -> 年份 -> 值
    project_dict = defaultdict(lambda: defaultdict(dict))

    for col in df_data.columns:
        year, type1, type2 = col
        value = row[col]
        if pd.notna(value):
            year_str = str(year).strip()
            key1 = type1.strip()
            key2 = type2.strip()
            if year_str in project_dict[key1][key2]:
                project_dict[key1][key2][year_str] += value
            else:
                project_dict[key1][key2][year_str] = value

    # 创建 Word 文档
    doc = Document()
    doc.add_heading(f"{grid_name} 投资项目汇总表", level=1)

    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '类型'
    hdr_cells[1].text = '项目名称'
    hdr_cells[2].text = '2025'
    hdr_cells[3].text = '2026'
    hdr_cells[4].text = '2027-2030'
    hdr_cells[5].text = '十五五合计'

    for type1 in type_order:
        if type1 in project_dict:
            for type2, year_vals in project_dict[type1].items():
                row_cells = table.add_row().cells
                row_cells[0].text = type1
                row_cells[1].text = type2
                row_cells[2].text = str(year_vals.get("2025", "/"))
                row_cells[3].text = str(year_vals.get("2026", "/"))
                row_cells[4].text = str(year_vals.get("2027-2030年", "/"))
                row_cells[5].text = str(year_vals.get("十五五合计", "/"))

    # 保存 Word 文件
    filename = f"{grid_name}.docx".replace("/", "_").replace("\\", "_")
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    print(f"已生成：{output_path}")
