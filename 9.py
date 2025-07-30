import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import re

os.makedirs("data_9", exist_ok=True)

def normalize(s):
    if pd.isna(s):
        return ""
    s = str(s).strip()
    s = s.replace("（", "(").replace("）", ")")
    s = s.replace("\xa0", "")
    s = s.replace(" ", "")
    s = re.sub(r"\s+", "", s)
    return s

def normalize_tuple(t):
    return tuple(normalize(x) for x in t)

# 读取数据
excel_path = "data/9.xlsx"
df = pd.read_excel(excel_path, skiprows=1)

# 选择和重命名列
df_selected = df[[
    '网格名称', '电压等级', '类型',
    '2024年装机容量', '2025年装机容量',
    '2026年装机容量', '2027年装机容量',
    '2028年装机容量', '2029年装机容量',
    '2030年装机容量'
]].copy()
df_selected.columns = ['名称', '电压等级', '类型', '装机2024', '装机2025', '装机2026', '装机2027', '装机2028', '装机2029', '装机2030']

# 拆分类型
df_selected[['类型1', '类型2']] = df_selected['类型'].str.split('-', n=1, expand=True)
df_selected['类型1'] = df_selected['类型1'].apply(normalize)
df_selected['类型2'] = df_selected['类型2'].apply(normalize)
df_selected['电压等级'] = df_selected['电压等级'].apply(normalize)

df_selected = df_selected[['名称', '电压等级', '类型1', '类型2', '装机2024', '装机2025', '装机2026', '装机2027', '装机2028', '装机2029', '装机2030']]

groups = df_selected.groupby('名称')

structure = [
    ('10（20）kV', '常规电源', '火电'),
    ('10（20）kV', '常规电源', '水电'),
    ('10（20）kV', '新能源', '风电'),
    ('10（20）kV', '新能源', '光伏'),
    ('10（20）kV', '新能源', '其中分布式'),
    ('10（20）kV', '新能源', '其他'),
    ('0.38kV及以下', '新能源', '风电'),
    ('0.38kV及以下', '新能源', '光伏'),
    ('0.38kV及以下', '新能源', '其中分布式'),
    ('0.38kV及以下', '新能源', '其他'),
    ('合计', '常规电源', '火电'),
    ('合计', '常规电源', '水电'),
    ('合计', '新能源', '风电'),
    ('合计', '新能源', '光伏'),
    ('合计', '新能源', '其中分布式'),
    ('合计', '新能源', '其他'),
]
normalized_structure = [normalize_tuple(t) for t in structure]

def set_font(cell, text):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

for name, group_df in groups:
    group_df = group_df.copy()
    group_df['电压等级'] = group_df['电压等级'].apply(normalize)
    group_df['类型1'] = group_df['类型1'].apply(normalize)
    group_df['类型2'] = group_df['类型2'].apply(normalize)

    doc = Document()
    p = doc.add_heading(f'表9  {name}电源装机容量及发电量  单位：MW，亿 kWh', level=1)
    run = p.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 创建表格，2行表头 + len(structure)行数据，10列
    table = doc.add_table(rows=2 + len(structure), cols=10)
    table.style = 'Table Grid'

    # 表头合并
    table.cell(0,0).merge(table.cell(1,0))  # 电压等级
    table.cell(0,1).merge(table.cell(1,2))  # 类型（跨2列）
    table.cell(0,3).merge(table.cell(0,9))  # 装机容量（跨7列）

    # 表头写入
    set_font(table.cell(0,0), '电压等级')
    set_font(table.cell(0,1), '类型')
    # 合并的类型第二行拆分2列，写细化指标
    set_font(table.cell(1,1), '类型1')
    set_font(table.cell(1,2), '类型2')

    # 装机容量子列写年份
    years = ['2024年', '2025年', '2026年', '2027年', '2028年', '2029年', '2030年']
    for i, y in enumerate(years):
        set_font(table.cell(1, 3+i), y)

    # 写数据
    for row_i, (dv_raw, t1_raw, t2_raw) in enumerate(structure):
        dv, t1, t2 = normalized_structure[row_i]
        row_cells = table.rows[row_i + 2].cells

        row_cells[0].text = dv_raw
        row_cells[1].text = t1_raw
        row_cells[2].text = t2_raw

        if dv != '合计':
            filtered = group_df[
                (group_df['电压等级'] == dv) &
                (group_df['类型1'] == t1) &
                (group_df['类型2'] == t2)
            ]
        else:
            # 合计行，跨电压等级求和，筛选类型1和类型2
            filtered = group_df[
                (group_df['类型1'] == t1) &
                (group_df['类型2'] == t2)
            ]

        if filtered.empty:
            # 没数据填“/”
            for i in range(7):
                row_cells[3+i].text = '/'
        else:
            sums = filtered[['装机2024', '装机2025', '装机2026', '装机2027', '装机2028', '装机2029', '装机2030']].sum()
            for i, y in enumerate(['装机2024', '装机2025', '装机2026', '装机2027', '装机2028', '装机2029', '装机2030']):
                val = sums[y]
                row_cells[3+i].text = f"{val:.2f}" if pd.notna(val) else '/'

    # 保存文件名安全处理
    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    output_word = f"data_9/{safe_name}.docx"
    doc.save(output_word)
    print(f"已保存：{output_word}")

print("所有文件已生成完成！")
