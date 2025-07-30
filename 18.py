import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_vmerge(cell, merge_type):
    """
    合并单元格（纵向）
    merge_type: 'restart'开始合并，'continue'继续合并
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vMerge = OxmlElement('w:vMerge')
    vMerge.set(qn('w:val'), merge_type)
    tcPr.append(vMerge)

# 读取数据
df_raw = pd.read_excel("data/18.xlsx", header=[0, 1])

# 提取“网格名称”与“类型”
df_raw[('网格名称', '')] = df_raw[('Unnamed: 0_level_0', '行标签')]
df_raw[('类型', '')] = df_raw[('Unnamed: 1_level_0', '行标签')].apply(lambda x: '新建' if '新建' in str(x) else ('改造' if '改造' in str(x) else '汇总'))

# 指标映射
indicator_map = {
    '电缆线': '电缆长度（km）'
}

# 输出路径
output_dir = "data_18"
os.makedirs(output_dir, exist_ok=True)

# 遍历网格分组
for grid_name, group_df in df_raw.groupby(('网格名称', '')):
    doc = Document()
    doc.add_heading(f'{grid_name} “十五五”配电网建设改造电缆线统计表', level=1)

    # 创建双行表头表格，先创建2行7列，数据行后续添加
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'

    hdr_cells1 = table.rows[0].cells
    hdr_cells2 = table.rows[1].cells

    # 前3列竖向合并两行
    for i, text in enumerate(["电压等级（kV）", "类型", "细化指标"]):
        hdr_cells1[i].text = text
        set_cell_vmerge(hdr_cells1[i], 'restart')
        set_cell_vmerge(hdr_cells2[i], 'continue')

    # 后4列，年份表头和“公用电网”
    years = ["2025", "2026", "2027-2030", "十五五合计"]
    for i, year in enumerate(years):
        hdr_cells1[i + 3].text = year
        hdr_cells2[i + 3].text = "公用电网"

    # 填充数据行
    # voltages = ['10kV', '0.38kV']
    voltages = ['10kV']
    types = ['汇总', '新建', '改造']

    for voltage in voltages:
        for t in types:
            row_data = group_df[group_df[('类型', '')] == t]
            if row_data.empty:
                continue
            for indicator_key, indicator_name in indicator_map.items():
                row = table.add_row().cells
                row[0].text = voltage
                row[1].text = t
                row[2].text = indicator_name

                for idx, year in enumerate(['2025年', '2026年', '2027年-2030年', '汇总']):
                    col = (year, indicator_key)
                    if col in row_data.columns:
                        val = row_data[col].values[0]
                        row[3 + idx].text = str(val) if pd.notna(val) else '\\'
                    else:
                        row[3 + idx].text = '\\'

    safe_name = grid_name.replace('/', '_').replace('\\', '_').replace(' ', '_')
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))
    print(f"已生成 {safe_name}.docx")
