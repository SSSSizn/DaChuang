import pandas as pd
from docx import Document
from collections import defaultdict
import os

# === 文件路径 ===
excel_path = "data/20.xlsx"
output_dir = "data_20"
os.makedirs(output_dir, exist_ok=True)

# === 读取 Excel，两层表头 ===
df = pd.read_excel(excel_path, header=[0, 1])

# === 拆分元数据与数据 ===
df_meta = df.iloc[:, :1]  # 第一列，网格名
df_data = df.iloc[:, 1:]  # 其余为数据列

# === 指标顺序和映射 ===
indicator_order = ["DTU", "FTU", "智能融合终端", "光缆", "ONU（套）"]
indicator_name_map = {
    "DTU": "DTU（台）",
    "FTU": "FTU（台）",
    "智能融合终端": "智能融合终端（台）",
    "光缆": "光缆（km）",
    "ONU（套）": "ONU(套)"
}

# === 年份映射 ===
year_map = {
    "2025年": "2025",
    "2026年": "2026",
    "2027年-2030年": "2027-2030",
    "求和项:": "十五五合计"
}

# === 每个网格一个 Word 文件 ===
for idx, row in df.iterrows():
    grid_name = row.iloc[0]

    # 构造指标 -> 年份 -> 值 结构
    data_dict = defaultdict(lambda: defaultdict(float))

    for col in df_data.columns:
        year_raw, ind_raw = col
        year = year_map.get(year_raw, year_raw)
        indicator = ind_raw.strip()

        if indicator in indicator_order:
            val = row[col]
            if pd.notna(val):
                data_dict[indicator][year] += val

    # 创建 Word 文档
    doc = Document()
    doc.add_heading(f"表20  {grid_name} “十五五”配电网自动化设备建设规模", level=1)

    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'

    hdr_cells1 = table.rows[0].cells
    hdr_cells1[0].text = "指标"
    hdr_cells1[1].text = "2025"
    hdr_cells1[2].text = "2026"
    hdr_cells1[3].text = "2027-2030"
    hdr_cells1[4].text = "十五五合计"

    hdr_cells2 = table.rows[1].cells
    hdr_cells2[0].text = ""
    for i in range(1, 5):
        hdr_cells2[i].text = "公用电网"

    # 填充数据从第三行开始
    for ind in indicator_order:
        row_cells = table.add_row().cells
        row_cells[0].text = indicator_name_map[ind]
        row_cells[1].text = str(data_dict[ind].get("2025", "\\"))
        row_cells[2].text = str(data_dict[ind].get("2026", "\\"))
        row_cells[3].text = str(data_dict[ind].get("2027-2030", "\\"))
        row_cells[4].text = str(data_dict[ind].get("十五五合计", "\\"))

    # 保存 Word
    filename = f"{grid_name}.docx".replace("/", "_").replace("\\", "_")
    doc.save(os.path.join(output_dir, filename))
    print(f"已生成：{filename}")