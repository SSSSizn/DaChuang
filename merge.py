from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import os

folder_list = [f"data_{i}" for i in [1, 2, 3, 4, 5, 6, 7, 9, 11, 13, 15, 16, 17, 18, 19, 20, 21]]
output_folder = "merged"
os.makedirs(output_folder, exist_ok=True)

grid_names = set()
for folder in folder_list:
    if not os.path.exists(folder):
        print(f"跳过不存在的文件夹: {folder}")
        continue
    for file in os.listdir(folder):
        if file.endswith(".docx"):
            grid_names.add(file)


def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)
a = 1
for grid_name in grid_names:
    print(f"{a}处理: {grid_name}")
    a = a + 1
    merged_doc = Document()

    for folder in folder_list:
        file_path = os.path.join(folder, grid_name)
        if os.path.exists(file_path):
            try:
                sub_doc = Document(file_path)
            except Exception as e:
                print(f"  无法读取 {file_path}：{e}")
                continue

            # 复制段落+表格内容
            for block in iter_block_items(sub_doc):
                if isinstance(block, Paragraph):
                    merged_doc.add_paragraph(block.text, style=block.style)
                elif isinstance(block, Table):
                    new_table = merged_doc.add_table(rows=0, cols=len(block.columns))
                    new_table.style = block.style
                    for row in block.rows:
                        new_row = new_table.add_row()
                        for idx, cell in enumerate(row.cells):
                            new_row.cells[idx].text = cell.text
                    merged_doc.add_paragraph()

    output_path = os.path.join(output_folder, grid_name)
    merged_doc.save(output_path)
    print(f"保存完成：{output_path}")

print("done")
