import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

input_path = "data/13.xlsx"  # 你源文件路径
output_dir = "data_13"
os.makedirs(output_dir, exist_ok=True)

# 读取Excel数据
df = pd.read_excel(input_path, header=1)

# 标准化列名去空格
df.columns = [c.strip() for c in df.columns]

# 需要的年份列
years = ['2024年', '2025年', '2026年', '2027年', '2028年', '2029年', '2030年']

# 需要的类型行（固定4种）
types = ['公共充电桩个数', '公共充电桩容量', '私人充电桩个数', '私人充电桩容量']

# 按网格名称分组
groups = df.groupby('网格名称')

for grid_name, group in groups:
    doc = Document()
    # 设置默认字体宋体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_heading(grid_name, level=1)

    # 创建表格，行=4（类型），列=1（类型列）+7（年份）
    table = doc.add_table(rows=len(types)+1, cols=1+len(years))
    table.style = 'Table Grid'

    # 写表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '类型'
    for i, y in enumerate(years):
        hdr_cells[i+1].text = y

    # 写数据行
    for i, t in enumerate(types):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = t

        # 找对应的行
        row_data = group[group['类型'] == t]

        if row_data.empty:
            # 没有这类型数据，全部填 /
            for j in range(len(years)):
                row_cells[j+1].text = '/'
        else:
            # 有数据，写对应年份数值
            for j, y in enumerate(years):
                val = row_data.iloc[0].get(y, '/')
                if pd.isna(val):
                    val = '/'
                else:
                    val = str(val)
                row_cells[j+1].text = val

    # 保存文件名安全处理
    safe_name = "".join(c for c in grid_name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))

print("所有网格充换电基础设施预测数据已导出完毕！")


