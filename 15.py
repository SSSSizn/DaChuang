import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import re

excel_path = "data/15.xlsx"
output_dir = "data_15"
os.makedirs(output_dir, exist_ok=True)

def normalize_text(s):
    """清洗文本，统一格式"""
    if pd.isna(s):
        return ""
    s = str(s).strip()
    # 全角转半角括号
    s = s.replace("（", "(").replace("）", ")")
    # 处理特殊空白字符
    s = s.replace("\xa0", "").replace(" ", "")
    # 去除所有空白字符（换行、制表符等）
    return re.sub(r"\s+", "", s)


def set_cell_font(cell, text):
    """设置 Word 表格单元格的字体格式（宋体、10 号）"""
    cell.text = ""  # 清空原有内容
    run = cell.paragraphs[0].add_run(text)
    # 设置英文字体
    run.font.name = 'SimSun'
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)


# 读取 Excel，skiprows=1 跳过第一行标题，从第二行开始作为表头和数据
df = pd.read_excel(excel_path, skiprows=1)

# 检查关键列是否存在，避免后续报错
required_columns = ["网格名称", "问题类型", "问题数量", "问题简述"]
missing_cols = [col for col in required_columns if col not in df.columns]
if missing_cols:
    raise ValueError(f"Excel 缺少必要列：{missing_cols}，请检查文件结构！")

# 网格名称分组，遍历每个网格的数据
grid_groups = df.groupby("网格名称")

for grid_name, group_data in grid_groups:
    safe_grid_name = "".join(
        c for c in grid_name if c.isalnum() or c in (' ', '_', '-')
    ).rstrip()
    # 拼接 Word 文档路径
    output_doc_path = os.path.join(output_dir, f"{safe_grid_name}.docx")

    doc = Document()

    # 添加标题，显示当前网格名称
    sub_title = doc.add_heading(f"表15 {grid_name}“十五五”中压配电网问题情况",  level=2)
    sub_title_run = sub_title.runs[0]
    sub_title_run.font.name = 'SimSun'
    sub_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    sub_title_run.font.size = Pt(13)
    sub_title_run.bold = False

    # 表格列：问题类型、问题数量、问题简述
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  # 带边框的表格样式

    # 表头内容
    headers = ["问题类型", "问题数量", "问题简述"]
    for col_idx, header in enumerate(headers):
        set_cell_font(table.cell(0, col_idx), header)

    # 先获取当前网格下的所有问题类型
    problem_types_in_excel = group_data["问题类型"].unique().tolist()

    # 遍历 Excel 中实际存在的问题类型
    for p_type in problem_types_in_excel:
        # 从分组数据中筛选当前问题类型的行
        filtered_data = group_data[group_data["问题类型"] == p_type]

        # 新增一行
        row_cells = table.add_row().cells

        # 填充「问题类型」
        set_cell_font(row_cells[0], p_type)

        # 填充「问题数量」：直接取 Excel 中的值
        if not filtered_data.empty:
            qty = filtered_data.iloc[0]["问题数量"]
            # 若问题数量是空值，显示空字符串；否则转成字符串填充
            set_cell_font(row_cells[1], str(qty) if not pd.isna(qty) else "")
        else:
            set_cell_font(row_cells[1], "")

        # 填充「问题简述」：直接取 Excel 中的值
        if not filtered_data.empty:
            desc = filtered_data.iloc[0]["问题简述"]
            # 清洗文本并填充
            set_cell_font(row_cells[2], normalize_text(desc))
        else:
            set_cell_font(row_cells[2], "")

    doc.save(output_doc_path)
    print(f"已生成：{output_doc_path}")

print("所有网格的 Word 文档已全部生成！")