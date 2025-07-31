import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import os

input_path = "data/6.xlsx"  # 修改为你的Excel路径
output_dir = "data_6"
os.makedirs(output_dir, exist_ok=True)

# 读取数据，假设无跳行，标题第一行
df = pd.read_excel(input_path,header=1)

# 清理列名空格和特殊字符
df.columns = [c.strip().replace(' ', '').replace(' ', '') for c in df.columns]  # 多替换了一个空格符号

# 需要的列（注意Excel里的列名和你的目标列名对齐）
cols_required = ['网格名称', '2020年', '2024年', '2025年', '2026年', '2027年', '2028年', '2029年', '2030年']

# 目标列顺序（加上你目标格式里多的两列）
target_cols = cols_required + ['十四五增速', '十五五增速']

# 判断源数据是否包含这两列
for col in ['十四五增速', '十五五增速']:
    if col not in df.columns:
        df[col] = '/'  # 不存在则新建列并赋值为缺失符号'/'

# 按网格名称分组
groups = df.groupby('网格名称')

for grid_name, group in groups:
    doc = Document()
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_heading(f"表6  {grid_name}全社会用电量预测表 单位：亿 kWh，%", level=1)

    # 取该网格的第一行数据（一般每个网格一条）
    row = group.iloc[0]

    # 创建表格 行数2（表头+数据），列数len(target_cols)
    table = doc.add_table(rows=2, cols=len(target_cols))
    table.style = 'Table Grid'

    # 写表头
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(target_cols):
        hdr_cells[i].text = col

    # 写数据行
    data_cells = table.rows[1].cells
    for i, col in enumerate(target_cols):
        val = row.get(col, '/')
        if pd.isna(val):
            val = '/'
        elif isinstance(val, float):
            val = f"{val:.2f}"
        else:
            val = str(val)
        data_cells[i].text = val

    # 保存，文件名安全处理
    safe_name = "".join(c for c in grid_name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    doc.save(f"{output_dir}/{safe_name}.docx")

print("所有网格全社会用电量预测已导出完毕！")