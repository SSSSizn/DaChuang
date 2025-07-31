import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# ======== 配置路径 ==========
input_path = "data/11.xlsx"  # 源文件路径
output_dir = "data_11"  # 输出目录
os.makedirs(output_dir, exist_ok=True)

# ======== 读取数据 ==========
df = pd.read_excel(input_path, header=1)

# 去除空格
df["电压等级"] = df["电压等级"].str.strip()
df["类型"] = df["类型"].str.strip()
df["网格名称"] = df["网格名称"].str.strip()

years = ["2024年", "2025年", "2026年", "2027年", "2028年", "2029年", "2030年"]

# 固定电压等级和类型组合，保证全量输出
fixed_voltages = ["10（20）kV", "0.38kV"]
fixed_types = ["电源侧", "电网侧", "用户侧"]

# 按网格名分组处理
for grid_name, group_df in df.groupby("网格名称"):
    # 汇总该网格下数据，按电压等级+类型分组求和
    summary = group_df.groupby(["电压等级", "类型"])[years].sum().reset_index()

    # 转成方便查询的字典
    data_dict = {}
    for _, row in summary.iterrows():
        key = (row["电压等级"], row["类型"])
        data_dict[key] = [row[y] if pd.notna(row[y]) else 0 for y in years]

    # 创建 Word 文档
    doc = Document()
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    doc.add_heading(f"表11  {grid_name}新型储能装机容量预测表  单位：MW", level=1)

    # 创建表格，2行表头 + 行数=电压等级*类型组合数，列数=3 + 年份数
    table = doc.add_table(rows=1 + len(fixed_voltages)*len(fixed_types), cols=2 + len(years))
    table.style = 'Table Grid'

    # 表头第一行
    hdr1 = table.rows[0].cells
    hdr1[0].text = "电压等级"
    hdr1[1].text = "类型"
    for i, y in enumerate(years):
        hdr1[2 + i].text = y

    # 表头第二行，年份列下写“公用电网”
    hdr2 = table.rows[1].cells
    hdr2[0].text = ""
    hdr2[1].text = ""
    hdr2[2].text = ""

    row_idx = 1
    for voltage in fixed_voltages:
        first_voltage_row = True
        for t in fixed_types:
            row_cells = table.rows[row_idx].cells
            # 电压等级第一行写，后续同电压等级行空着
            if first_voltage_row:
                row_cells[0].text = voltage
                first_voltage_row = False
            else:
                row_cells[0].text = ""
            row_cells[1].text = t

            # 填值，没有则填“\”
            values = data_dict.get((voltage, t), ["\\"]*len(years))
            for i, v in enumerate(values):
                # 保留两位小数，如果是数字
                if v == "\\":
                    row_cells[2 + i].text = "\\"
                else:
                    try:
                        row_cells[2 + i].text = f"{float(v):.2f}"
                    except:
                        row_cells[2 + i].text = "\\"
            row_idx += 1

    # 保存文件，文件名里替换掉可能的斜杠等特殊字符
    safe_name = grid_name.replace("/", "-").replace("\\", "-")
    doc.save(os.path.join(output_dir, f"{safe_name}.docx"))

print("所有网格的储能装机容量预测表已导出完成！")