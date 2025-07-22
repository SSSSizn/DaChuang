import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import re

# 自动创建输出目录
os.makedirs("data_2", exist_ok=True)

# 字符清洗函数
def normalize(s):
    s = str(s).strip()
    s = s.replace("（", "(").replace("）", ")")  # 全角转半角括号
    s = s.replace("\xa0", "")  # 非断行空格（U+00A0）
    s = s.replace(" ", "")  # 去普通空格
    s = re.sub(r"\s+", "", s)  # 去掉所有空白字符
    return s

def normalize_tuple(t):
    return tuple(normalize(x) for x in t)

# 读取 Excel 数据
excel_path = "data/2.xlsx"
df = pd.read_excel(excel_path, skiprows=1)

# 选择和重命名列
df_selected = df[[
    '网格名称', '电压等级', '类型',
    '2020年装机容量（MW）', '2024年装机容量（MW）',
    '2020年发电量（亿kWh）', '2024年发电量（亿kWh）'
]].copy()
df_selected.columns = ['名称', '电压等级', '类型', '装机2020', '装机2024', '发电2020', '发电2024']

# 拆分类型并清洗
df_selected[['类型1', '类型2']] = df_selected['类型'].str.split('-', n=1, expand=True)
df_selected['类型1'] = df_selected['类型1'].apply(normalize)
df_selected['类型2'] = df_selected['类型2'].apply(normalize)
df_selected['电压等级'] = df_selected['电压等级'].apply(normalize)

# 重整列顺序
df_selected = df_selected[['名称', '电压等级', '类型1', '类型2', '装机2020', '装机2024', '发电2020', '发电2024']]
groups = df_selected.groupby('名称')

# 表格字体设置函数
def set_font(cell, text):
    run = cell.paragraphs[0].add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

# 定义结构并清洗
structure = [
    ('10（20）kV', '常规电源', '水电'),
    ('10（20）kV', '常规电源', '火电'),
    ('10（20）kV', '新能源', '风电'),
    ('10（20）kV', '新能源', '光伏'),
    ('10（20）kV', '新能源', '其中分布式'),
    ('10（20）kV', '新能源', '其他'),
    ('0.38kV 及以下', '新能源', '风电'),
    ('0.38kV 及以下', '新能源', '光伏'),
    ('0.38kV 及以下', '新能源', '其中分布式'),
    ('0.38kV 及以下', '新能源', '其他'),
    ('合计', '常规电源', '火电'),
    ('合计', '常规电源', '水电'),
    ('合计', '新能源', '风电'),
    ('合计', '新能源', '光伏'),
    ('合计', '新能源', '其中分布式'),
    ('合计', '新能源', '其他'),
]
normalized_structure = [normalize_tuple(t) for t in structure]

# 遍历每个网格
for name, group_df in groups:
    group_df = group_df.copy()
    group_df['电压等级'] = group_df['电压等级'].apply(normalize)
    group_df['类型1'] = group_df['类型1'].apply(normalize)
    group_df['类型2'] = group_df['类型2'].apply(normalize)

    print(f"\n--- 网格：{name} ---")
    group_keys = set(zip(group_df['电压等级'], group_df['类型1'], group_df['类型2']))
    for key in normalized_structure:
        if key not in group_keys and key[0] != '合计':
            print(f"无法匹配：{key} 不存在于数据中")

    doc = Document()
    p = doc.add_heading(f'表2  {name}电源装机容量及发电量  单位：MW，亿 kWh', level=1)
    run = p.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'

    set_font(table.cell(0, 0), '电压等级')
    set_font(table.cell(0, 1), '类型')
    set_font(table.cell(0, 3), '装机容量')
    set_font(table.cell(0, 5), '发电量')

    table.cell(0, 0).merge(table.cell(1, 0))
    table.cell(0, 1).merge(table.cell(1, 2))
    table.cell(0, 3).merge(table.cell(0, 4))
    table.cell(0, 5).merge(table.cell(0, 6))

    set_font(table.cell(1, 3), '2020年')
    set_font(table.cell(1, 4), '2024年')
    set_font(table.cell(1, 5), '2020年')
    set_font(table.cell(1, 6), '2024年')

    # 写入数据行
    for original_def, norm_def in zip(structure, normalized_structure):
        dv_raw, t1_raw, t2_raw = original_def
        dv, t1, t2 = norm_def

        if dv != '合计':
            row_data = group_df[
                (group_df['电压等级'] == dv) &
                (group_df['类型1'] == t1) &
                (group_df['类型2'] == t2)
                ]
        else:
            row_data = group_df[
                (group_df['类型1'] == t1) &
                (group_df['类型2'] == t2)
                ]

        if row_data.empty:
            continue  # 跳过无数据行

        cells = table.add_row().cells
        cells[0].text = dv_raw
        cells[1].text = t1_raw
        cells[2].text = t2_raw

        agg = row_data[['装机2020', '装机2024', '发电2020', '发电2024']].sum()
        cells[3].text = '0' if pd.isna(agg['装机2020']) else str(agg['装机2020'])
        cells[4].text = '0' if pd.isna(agg['装机2024']) else str(agg['装机2024'])
        cells[5].text = '0' if pd.isna(agg['发电2020']) else str(agg['发电2020'])
        cells[6].text = '0' if pd.isna(agg['发电2024']) else str(agg['发电2024'])

    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    output_word = f"data_2/{safe_name}.docx"
    doc.save(output_word)
    print(f"已保存：{output_word}")
