import pandas as pd
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

# 读取 Excel 文件
excel_path = 'data/19.xlsx'
output_dir = 'data_19'
os.makedirs(output_dir, exist_ok=True)

# 读取表格
df_raw = pd.read_excel(excel_path, header=[0, 1])
df_raw = df_raw.dropna(how='all')  # 去除空行

# 拆解列名结构
df_raw.columns = [tuple(map(str, col)) for col in df_raw.columns]
df_raw.columns = pd.MultiIndex.from_tuples(df_raw.columns)

# 填充第一列行标签
df_raw.iloc[:, 0] = df_raw.iloc[:, 0].ffill()

# 提取主要列
df = df_raw.copy()
df.columns.names = [None, None]
df.rename(columns={df.columns[0]: ('网格名称', '')}, inplace=True)
df.rename(columns={df.columns[1]: ('类型', '')}, inplace=True)

# 设置指标映射关系
indicator_map = {
    '配变': ('配变（台）', '台'),
    '配变容量': ('容量（kVA）', 'kVA'),
    '环网箱': ('环网箱（座）', '座'),
    '环网室': ('环网室（座）', '座'),
    '柱上开关': ('柱上开关（台）', '台'),
    '配变汇总': ('配变（台）', '台'),
    '配变容量汇总': ('容量（kVA）', 'kVA'),
    '环网箱汇总': ('环网箱（座）', '座'),
    '环网室汇总': ('环网室（座）', '座'),
    '柱上开关汇总': ('柱上开关（台）', '台'),
}

years = ['2025年', '2026年', '2027年-2030年', '求和项']
voltage_level = '10'

# 将前两列提出来作为普通列
df.columns = pd.MultiIndex.from_tuples(df.columns)
df = df.copy()
df.columns.names = [None, None]

# 提取“网格名称”和“类型”作为普通列
df['网格名称'] = df.iloc[:, 0]
df['类型'] = df.iloc[:, 1]

# 去除原始列
df = df.drop(columns=[df.columns[0], df.columns[1]])

# 按网格名称分组
for grid_name, group_df in df.groupby(('网格名称', '')):
    doc = Document()
    table = doc.add_table(rows=2, cols=7)
    table.style = 'Table Grid'

    # 表头第一行
    headers = ['电压等级（kV）', '类型', '细化指标', '2025', '2026', '2027-2030', '十五五合计']
    for i, text in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = text
        cell.paragraphs[0].runs[0].font.size = Pt(10)

    # 表头第二行
    for i in range(3):
        table.cell(1, i).text = ''
    for i in range(3, 7):
        table.cell(1, i).text = '公用电网'

    # 数据填充
    for _, row in group_df.iterrows():
        row_type = row[('类型', '')].strip()
        for col_key in df.columns[2:]:
            year, indicator = col_key
            if indicator not in indicator_map or year not in years:
                continue

            indicator_name, unit = indicator_map[indicator]
            value = row.get((year, indicator), '\\')

            # 查找或新增该指标对应的行
            existing_row = None
            for r in table.rows[2:]:
                if (r.cells[1].text == row_type) and (r.cells[2].text == indicator_name):
                    existing_row = r
                    break

            if not existing_row:
                r = table.add_row()
                r.cells[0].text = voltage_level
                r.cells[1].text = row_type
                r.cells[2].text = indicator_name
                r.cells[3].text = '\\'
                r.cells[4].text = '\\'
                r.cells[5].text = '\\'
                r.cells[6].text = '\\'
                existing_row = r

            year_idx = years.index(year)
            existing_row.cells[3 + year_idx].text = str(value if pd.notna(value) else '\\')

    # 设置字体为中文宋体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(10)

    # 保存文档
    filename = os.path.join(output_dir, f'{grid_name}.docx')
    doc.save(filename)
    print(f"已保存: {filename}")
