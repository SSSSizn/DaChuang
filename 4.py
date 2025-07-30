import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import os
import re

# 自动创建输出目录
os.makedirs("data_4", exist_ok=True)

# 字符清洗函数
def normalize(s):
    s = str(s).strip()
    s = s.replace("（", "(").replace("）", ")")  # 全角转半角括号
    s = s.replace("\xa0", "")  # 非断行空格（U+00A0）
    s = s.replace(" ", "")  # 去普通空格
    s = re.sub(r"\s+", "", s)  # 去掉所有空白字符
    return s

def normalize_tuple(t):
    return tuple(normalize(x) for x in t)

# 读取 Excel 数据
excel_path = "data/4.xlsx"
df = pd.read_excel(excel_path, skiprows=1)

# 选择和重命名列
df_selected = df[[
    '网格名称', '电压等级', '类型',
    '2024年','2025年','2026年','2027年','2028年','2029年','2030年'
]].copy()
df_selected.columns = ['名称', '电压等级', '类型', '2024年','2025年','2026年','2027年','2028年','2029年','2030年']

# 清洗字段
for col in ['电压等级', '类型']:
    df_selected[col] = df_selected[col].apply(normalize)


groups = df_selected.groupby('名称')

# 表格字体设置函数
def set_font(cell, text):
    run = cell.paragraphs[0].add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(10)

# 定义结构并清洗
structure = [
    ('10（20）kV', '电源侧'),
    ('10（20）kV', '电网侧'),
    ('10（20）kV', '用户侧'),
    ('0.38kV', '电源侧'),
    ('0.38kV', '电网侧'),
    ('0.38kV', '用户侧'),
]
normalized_structure = [normalize_tuple(t) for t in structure]

# 遍历每个网格
for name, group_df in groups:
    doc = Document()
    p = doc.add_heading(f'表4  {name}新型储能装机容量      单位：MW', level=1)
    run = p.runs[0]
    run.font.name = '宋体'
    run.font.size = Pt(10)
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 表头
    table = doc.add_table(rows=1, cols=9)
    table.style = 'Table Grid'
    headers = ['电压等级', '类型', '2024年','2025年','2026年','2027年','2028年','2029年','2030年']
    for col_idx, header in enumerate(headers):
        set_font(table.cell(0, col_idx), header)

    for (dv_raw, t_raw), (dv, t) in zip(structure, normalized_structure):
        row_data = group_df[
            (group_df['电压等级'] == dv) &
            (group_df['类型'] == t)
        ]

        row = table.add_row().cells
        row[0].text = dv_raw
        row[1].text = t_raw

        if row_data.empty:
            for i in range(2, 9):
                row[i].text = '/'
        else:
            sums = row_data[['2024年','2025年','2026年','2027年','2028年','2029年','2030年']].sum()
            for i, y in enumerate(['2024年','2025年','2026年','2027年','2028年','2029年','2030年']):
                val = sums[y]
                row[i + 2].text = '/' if pd.isna(val) else str(round(val, 2))

    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '_', '-')).rstrip()
    output_path = f"data_4/{safe_name}.docx"
    doc.save(output_path)
    print(f"已生成：{output_path}")